#!/usr/bin/env python3
"""
TRIZ 报告 Markdown -> Word 转换脚本
功能：封面页、三级目录、章节分页、格式规范
"""

import re
import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_cover_page(doc, title, subtitle, date):
    """添加封面页"""
    # 添加空行使内容居中
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 40)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xF6, 0xAD, 0x55)
    
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # 分页
    doc.add_page_break()


def add_toc(doc):
    """添加目录页"""
    # 目录标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    doc.add_paragraph()  # 空行
    
    # 添加 TOC 域代码
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)
    
    run = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)
    
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar)
    
    # 添加占位文本
    run = paragraph.add_run('（请在 Word 中右键点击此处，选择"更新域"以生成目录）')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    run.font.italic = True
    
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar)
    
    # 分页
    doc.add_page_break()


def parse_markdown(md_content):
    """解析 Markdown 内容为结构化数据"""
    lines = md_content.split('\n')
    sections = []
    current_section = None
    current_content = []
    
    for line in lines:
        # 检测标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            # 保存当前section
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)
            
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = {
                'level': level,
                'title': title,
                'content': ''
            }
            current_content = []
        else:
            current_content.append(line)
    
    # 保存最后一个section
    if current_section:
        current_section['content'] = '\n'.join(current_content)
        sections.append(current_section)
    
    return sections


def parse_table(content):
    """解析 Markdown 表格"""
    lines = content.strip().split('\n')
    table_lines = [l for l in lines if '|' in l and l.strip().startswith('|')]
    
    if len(table_lines) < 2:
        return None
    
    # 解析表头
    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    
    # 跳过分隔行
    rows = []
    for line in table_lines[2:]:  # 跳过表头和分隔行
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    return {'headers': headers, 'rows': rows}


def add_table_to_doc(doc, table_data):
    """添加表格到文档"""
    if not table_data:
        return
    
    headers = table_data['headers']
    rows = table_data['rows']
    
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1A365D')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text
                # 交替行颜色
                if row_idx % 2 == 0:
                    set_cell_shading(cell, 'F7FAFC')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    doc.add_paragraph()  # 表格后空行


def add_content_to_doc(doc, content, is_first_level=False):
    """添加内容到文档"""
    lines = content.split('\n')
    in_table = False
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # 跳过空行
        if not stripped:
            if in_table and table_lines:
                table_data = parse_table('\n'.join(table_lines))
                if table_data:
                    add_table_to_doc(doc, table_data)
                table_lines = []
                in_table = False
            continue
        
        # 检测表格
        if '|' in stripped and stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        
        if in_table and table_lines:
            table_data = parse_table('\n'.join(table_lines))
            if table_data:
                add_table_to_doc(doc, table_data)
            table_lines = []
            in_table = False
        
        # 普通文本
        # 清理 Markdown 格式
        text = stripped
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 去除加粗
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # 去除斜体
        text = re.sub(r'`(.+?)`', r'\1', text)  # 去除代码
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # 去除链接
        
        # 跳过代码块标记
        if text.startswith('```'):
            continue
        
        # 检测列表项
        list_match = re.match(r'^[-*]\s+(.+)$', text)
        numbered_match = re.match(r'^\d+\.\s+(.+)$', text)
        
        if list_match or numbered_match:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(list_match.group(1) if list_match else numbered_match.group(1))
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5


def md_to_word(md_file, output_file, title=None, date=None):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # 获取标题
    if not title:
        # 从文件名提取
        basename = os.path.basename(md_file)
        title = basename.replace('_完整版.md', '').replace('TRIZ_报告_', '').replace('TRIZ_创新分析报告_', '')
        title = f'TRIZ 创新分析报告\n{title}'
    
    if not date:
        date = datetime.now().strftime('%Y年%m月%d日')
    
    # 添加封面页
    add_cover_page(doc, title, 'TRIZ 系统化创新分析报告', date)
    
    # 添加目录页
    add_toc(doc)
    
    # 解析并添加内容
    sections = parse_markdown(md_content)
    
    for i, section in enumerate(sections):
        level = section['level']
        title = section['title']
        content = section['content']
        
        # 一级标题前分页
        if level == 1 and i > 0:
            doc.add_page_break()
        
        # 添加标题
        heading = doc.add_heading(title, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        # 添加内容
        add_content_to_doc(doc, content, is_first_level=(level == 1))
    
    # 保存文档
    doc.save(output_file)
    print(f'✅ Word 文档已生成: {output_file}')
    
    return output_file


def main():
    if len(sys.argv) < 2:
        print('用法: python triz_md_to_docx.py <input.md> [output.docx] [--title "标题"] [--date "日期"]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # 解析参数
    output_file = None
    title = None
    date = None
    
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '--title' and i + 1 < len(sys.argv):
            title = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == '--date' and i + 1 < len(sys.argv):
            date = sys.argv[i + 1]
            i += 2
        elif not output_file:
            output_file = sys.argv[i]
            i += 1
        else:
            i += 1
    
    if not output_file:
        output_file = input_file.replace('.md', '.docx')
    
    md_to_word(input_file, output_file, title, date)


if __name__ == '__main__':
    main()
